from .base import BaseExtractionStrategy, ExtractionResult
import logging
from docx import Document
import io

logger = logging.getLogger(__name__)

class DocxStrategy(BaseExtractionStrategy):
    """Extracts text from Word Documents (.docx)."""

    def extract(self, file_stream, language: str = 'eng') -> ExtractionResult:
        logger.info("Starting Docx extraction.")
        try:
            # Ensure at start
            if hasattr(file_stream, 'seek'):
                file_stream.seek(0)
            
            # python-docx can open file objects directly
            # Note: python-docx loads the entire file into memory. For very large DOCX files,
            # this may cause memory issues. Unfortunately, the library doesn't support streaming.
            
            # Check file size if possible
            if hasattr(file_stream, 'seek') and hasattr(file_stream, 'tell'):
                file_stream.seek(0, 2)  # Seek to end
                file_size_mb = file_stream.tell() / (1024 * 1024)
                file_stream.seek(0)  # Reset
                
                if file_size_mb > 50:
                    logger.warning(f"Large DOCX file detected: {file_size_mb:.1f} MB. This may use significant memory.")
            
            doc = Document(file_stream)
            
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
                
            text_content = "\n".join(full_text)
            
            # Docx doesn't map strictly to "pages" in a print sense without rendering, 
            # so we treat it as a continuous document or paragraphs?
            # For compatibility with our result structure, let's treat it as one single "page" 
            # or maybe split by hard breaks if we wanted.
            # Let's verify if python-docx exposes pages. No, it doesn't.
            
            pages = [text_content]

            logger.info(f"Docx extraction successful. extracted {len(doc.paragraphs)} paragraphs.")
            
            return ExtractionResult(
                full_text=text_content,
                pages=pages,
                metadata={"method": "docx", "page_count": 1, "paragraph_count": len(doc.paragraphs)}
            )

        except Exception as e:
            logger.error(f"Docx extraction failed: {e}")
            raise Exception(f"Failed to extract text from Docx: {str(e)}")
